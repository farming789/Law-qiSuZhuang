import streamlit as st
import os
from langchain_community.chat_models.tongyi import ChatTongyi
from langchain.prompts import ChatPromptTemplate
from langchain.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field
from typing import List, Optional
import docx
import io

# --- 1. 定义数据结构 (Pydantic Models) ---
# 定义清晰的数据模型，用于规范LLM的输出，确保数据格式的统一和稳定

class PartyInfo(BaseModel):
    """用于存储原告或被告信息的模型"""
    name: Optional[str] = Field(default=None, description="姓名")
    gender: Optional[str] = Field(default=None, description="性别")
    ethnicity: Optional[str] = Field(default=None, description="民族")
    dob: Optional[str] = Field(default=None, description="出生年月日")
    address: Optional[str] = Field(default=None, description="住址")
    id_card: Optional[str] = Field(default=None, description="公民身份证号码")
    contact: Optional[str] = Field(default=None, description="联系电话")

class Lawsuit(BaseModel):
    """起诉状的完整数据结构模型"""
    plaintiff: PartyInfo = Field(description="原告信息")
    defendant: PartyInfo = Field(description="被告信息")
    claims: str = Field(description="诉讼请求，保持原始的多点格式，作为一个单一的字符串")
    facts_and_reasons: str = Field(description="事实与理由的完整陈述")
    court_name: Optional[str] = Field(default=None, description="提交的法院名称")
    date: Optional[str] = Field(default=None, description="起诉状的日期")


# --- 2. Langchain 信息提取模块 ---

@st.cache_data
def extract_lawsuit_data(file_path: str, api_key: str) -> Optional[Lawsuit]:
    """
    使用Langchain从上传的Word文档中提取信息。
    - file_path: 上传的原始起诉状文件路径。
    - api_key: DashScope API Key.
    - 返回: 一个填充了数据的Lawsuit对象，如果失败则返回None。
    """
    try:
        # 1. 加载文档（仅支持 .docx），使用 python-docx 读取，避免对 LibreOffice 的依赖
        if not file_path.lower().endswith(".docx"):
            raise ValueError("仅支持 .docx 文件，请将 .doc 文件另存为 .docx 后再上传。")
        docx_doc = docx.Document(file_path)
        parts = []
        for para in docx_doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in docx_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        parts.append(cell_text)
        content = "\n".join(parts)

        # 2. 设置输出解析器
        parser = PydanticOutputParser(pydantic_object=Lawsuit)

        # 3. 创建提示词模板
        prompt_template = """
        你是一个专业的法律文书助手。请从以下起诉状文本中提取所有关键信息。
        请严格按照以下JSON格式指令输出结果，如果某些信息在文本中不存在，请将其值留空(null)。

        文本内容:
        {document_text}

        格式指令:
        {format_instructions}
        """
        prompt = ChatPromptTemplate.from_template(
            template=prompt_template,
            partial_variables={"format_instructions": parser.get_format_instructions()}
        )

        # 4. 初始化LLM模型
        model = ChatTongyi(dashscope_api_key=api_key, model="qwen3-235b-a22b-thinking-2507", temperature=0)

        # 5. 创建并执行链
        chain = prompt | model | parser
        result = chain.invoke({"document_text": content})
        return result
    except Exception as e:
        st.error(f"信息提取失败: {e}")
        return None

# --- 3. Word 文档生成模块 ---

def generate_docx(data: Lawsuit, template_path: str) -> Optional[bytes]:
    """
    将提取并审核过的数据填充到Word模板中。
    - data: 包含最终数据的Lawsuit对象。
    - template_path: 带有占位符的.docx模板文件路径。
    - 返回: 生成的Word文档的字节流，如果失败则返回None。
    """
    try:
        doc = docx.Document(template_path)
        
        # 创建一个字典，方便替换
        # 注意：这里的key需要与你在template.docx中设置的占位符完全一致
        replacements = {
            "{{plaintiff_name}}": data.plaintiff.name or "",
            "{{plaintiff_gender}}": data.plaintiff.gender or "",
            "{{plaintiff_ethnicity}}": data.plaintiff.ethnicity or "",
            "{{plaintiff_dob}}": data.plaintiff.dob or "",
            "{{plaintiff_address}}": data.plaintiff.address or "",
            "{{plaintiff_id_card}}": data.plaintiff.id_card or "",
            "{{plaintiff_contact}}": data.plaintiff.contact or "",
            "{{defendant_name}}": data.defendant.name or "",
            "{{defendant_gender}}": data.defendant.gender or "",
            "{{defendant_ethnicity}}": data.defendant.ethnicity or "",
            "{{defendant_dob}}": data.defendant.dob or "",
            "{{defendant_address}}": data.defendant.address or "",
            "{{defendant_id_card}}": data.defendant.id_card or "",
            "{{defendant_contact}}": data.defendant.contact or "",
            "{{claims}}": data.claims or "",
            "{{facts_and_reasons}}": data.facts_and_reasons or "",
            "{{court_name}}": data.court_name or "",
            "{{date}}": data.date or "",
        }

        # 替换段落中的占位符
        for para in doc.paragraphs:
            for key, value in replacements.items():
                if key in para.text:
                    # 使用run来替换，可以保留部分格式
                    # 这是一个简化的替换，复杂的格式可能需要更精细的处理
                    inline = para.runs
                    for i in range(len(inline)):
                        if key in inline[i].text:
                            text = inline[i].text.replace(key, value)
                            inline[i].text = text

        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in para.text:
                                inline = para.runs
                                for i in range(len(inline)):
                                    if key in inline[i].text:
                                        text = inline[i].text.replace(key, value)
                                        inline[i].text = text
        
        # 将文档保存到内存中的字节流
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()

    except Exception as e:
        st.error(f"文档生成失败: {e}")
        return None


# --- 4. Streamlit 用户界面 ---

st.set_page_config(page_title="起诉状格式转换工具", layout="wide")
st.title("⚖️ 起诉状智能格式转换工具")

# 获取通义千问 API Key
api_key = os.getenv("DASHSCOPE_API_KEY")
if not api_key:
    api_key = st.text_input("请输入你的 通义千问 API Key (DashScope):", type="password")


# 文件上传
uploaded_file = st.file_uploader(
    "上传原始起诉状 (.docx)",
    type=['docx']
)

# Session state 用于存储提取的数据
if 'extracted_data' not in st.session_state:
    st.session_state.extracted_data = None

if uploaded_file is not None and api_key:
    # 将上传的文件保存到临时位置
    with open("temp_uploaded_file.docx", "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    if st.button("开始提取信息"):
        with st.spinner("正在调用大模型分析文档，请稍候..."):
            st.session_state.extracted_data = extract_lawsuit_data("temp_uploaded_file.docx", api_key)
        if st.session_state.extracted_data:
            st.success("信息提取成功！请在下方表单中审核和修改。")
        else:
            st.error("无法从文档中提取信息，请检查文件内容或API Key。")

# 如果数据已提取，显示审核表单
if st.session_state.extracted_data:
    data = st.session_state.extracted_data
    
    with st.form("review_form"):
        st.header("审核与修改")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("原告信息")
            p_name = st.text_input("姓名", value=data.plaintiff.name, key="p_name")
            p_gender = st.text_input("性别", value=data.plaintiff.gender, key="p_gender")
            p_ethnicity = st.text_input("民族", value=data.plaintiff.ethnicity, key="p_ethnicity")
            p_dob = st.text_input("出生年月日", value=data.plaintiff.dob, key="p_dob")
            p_address = st.text_area("住址", value=data.plaintiff.address, key="p_address")
            p_id = st.text_input("身份证号", value=data.plaintiff.id_card, key="p_id")
            p_contact = st.text_input("联系电话", value=data.plaintiff.contact, key="p_contact")

        with col2:
            st.subheader("被告信息")
            d_name = st.text_input("姓名", value=data.defendant.name, key="d_name")
            d_gender = st.text_input("性别", value=data.defendant.gender, key="d_gender")
            d_ethnicity = st.text_input("民族", value=data.defendant.ethnicity, key="d_ethnicity")
            d_dob = st.text_input("出生年月日", value=data.defendant.dob, key="d_dob")
            d_address = st.text_area("住址", value=data.defendant.address, key="d_address")
            d_id = st.text_input("身份证号", value=data.defendant.id_card, key="d_id")
            d_contact = st.text_input("联系电话", value=data.defendant.contact, key="d_contact")

        st.subheader("诉讼请求")
        claims_text = st.text_area("诉讼请求", value=data.claims, height=200, key="claims")

        st.subheader("事实与理由")
        facts_text = st.text_area("事实与理由", value=data.facts_and_reasons, height=400, key="facts")
        
        st.subheader("其他信息")
        court_name_text = st.text_input("法院名称", value=data.court_name, key="court_name")
        date_text = st.text_input("日期", value=data.date, key="date")

        submitted = st.form_submit_button("确认信息并生成文档")

        if submitted:
            # 收集表单中的最新数据
            final_data = Lawsuit(
                plaintiff=PartyInfo(
                    name=st.session_state.p_name, gender=st.session_state.p_gender, ethnicity=st.session_state.p_ethnicity,
                    dob=st.session_state.p_dob, address=st.session_state.p_address, id_card=st.session_state.p_id,
                    contact=st.session_state.p_contact
                ),
                defendant=PartyInfo(
                    name=st.session_state.d_name, gender=st.session_state.d_gender, ethnicity=st.session_state.d_ethnicity,
                    dob=st.session_state.d_dob, address=st.session_state.d_address, id_card=st.session_state.d_id,
                    contact=st.session_state.d_contact
                ),
                claims=st.session_state.claims,
                facts_and_reasons=st.session_state.facts,
                court_name=st.session_state.court_name,
                date=st.session_state.date
            )
            
            with st.spinner("正在生成标准格式的Word文档..."):
                # 检查模板文件是否存在
                if not os.path.exists("template.docx"):
                    st.error("错误：未找到模板文件 `template.docx`。请确保它与 `app.py` 在同一目录下。")
                else:
                    generated_doc_bytes = generate_docx(final_data, "template.docx")
            
                    if generated_doc_bytes:
                        st.success("文档生成成功！")
                        st.download_button(
                            label="下载高标准格式起诉状.docx",
                            data=generated_doc_bytes,
                            file_name="高标准格式的起诉状_完成.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

# 清理临时文件
if os.path.exists("temp_uploaded_file.docx"):
    os.remove("temp_uploaded_file.docx")

st.info("说明：请先在上方输入框中提供您的通义千问 (DashScope) API Key，然后上传原始起诉状文件，点击“开始提取信息”按钮。AI提取信息后，您可以在表单中进行修改，最后点击“确认信息并生成文档”来下载最终文件。")
